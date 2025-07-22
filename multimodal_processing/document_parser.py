"""
Multimodal Document Parser for SAM
Ingests and parses PDFs, DOCX, Markdown, and HTML containing tables, text, code snippets, and image metadata.

Sprint 4 Task 1: Multimodal Ingestion Pipeline
"""

import logging
import json
import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass
import hashlib
from datetime import datetime

logger = logging.getLogger(__name__)

@dataclass
class MultimodalContent:
    """Represents a piece of multimodal content."""
    content_type: str  # 'text', 'code', 'table', 'image'
    content: Union[str, List[List[str]], Dict[str, Any]]
    metadata: Dict[str, Any]
    source_location: Optional[str] = None

@dataclass
class ParsedDocument:
    """Represents a fully parsed multimodal document."""
    document_id: str
    source_file: str
    content_blocks: List[MultimodalContent]
    document_metadata: Dict[str, Any]
    parsing_stats: Dict[str, Any]

class MultimodalDocumentParser:
    """
    Parses various document formats and extracts multimodal content.
    """
    
    def __init__(self):
        """Initialize the multimodal document parser."""
        self.supported_formats = {'.pdf', '.docx', '.md', '.html', '.htm', '.txt', '.py', '.js', '.java', '.cpp', '.c'}
        self.parsers = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.md': self._parse_markdown,
            '.html': self._parse_html,
            '.htm': self._parse_html,
            '.txt': self._parse_text,
            '.py': self._parse_code_file,
            '.js': self._parse_code_file,
            '.java': self._parse_code_file,
            '.cpp': self._parse_code_file,
            '.c': self._parse_code_file
        }
        
        logger.info(f"Multimodal document parser initialized")
        logger.info(f"Supported formats: {', '.join(self.supported_formats)}")
    
    def parse_document(self, file_path: Union[str, Path]) -> Optional[ParsedDocument]:
        """
        Parse a document and extract multimodal content.
        
        Args:
            file_path: Path to the document to parse
            
        Returns:
            ParsedDocument object or None if parsing failed
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            logger.warning(f"Unsupported file format: {file_extension}")
            return None
        
        try:
            logger.info(f"Parsing document: {file_path}")
            
            # Generate document ID
            document_id = self._generate_document_id(file_path)
            
            # Parse using appropriate parser
            parser_func = self.parsers[file_extension]
            content_blocks = parser_func(file_path)
            
            # Create document metadata
            document_metadata = {
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'file_extension': file_extension,
                'parsed_timestamp': datetime.now().isoformat(),
                'content_block_count': len(content_blocks)
            }
            
            # Calculate parsing stats
            parsing_stats = self._calculate_parsing_stats(content_blocks)
            
            parsed_doc = ParsedDocument(
                document_id=document_id,
                source_file=str(file_path),
                content_blocks=content_blocks,
                document_metadata=document_metadata,
                parsing_stats=parsing_stats
            )
            
            logger.info(f"Successfully parsed {file_path}: {len(content_blocks)} content blocks")
            return parsed_doc
            
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {e}")
            return None
    
    def _generate_document_id(self, file_path: Path) -> str:
        """Generate a unique document ID."""
        content = f"{file_path.name}_{file_path.stat().st_mtime}_{file_path.stat().st_size}"
        return f"doc_{hashlib.md5(content.encode()).hexdigest()[:12]}"
    
    def _parse_pdf(self, file_path: Path) -> List[MultimodalContent]:
        """Parse PDF documents."""
        content_blocks = []
        
        try:
            # Try to import PDF parsing libraries
            try:
                import PyPDF2
                pdf_parser = 'PyPDF2'
            except ImportError:
                try:
                    import pdfplumber
                    pdf_parser = 'pdfplumber'
                except ImportError:
                    logger.warning("No PDF parsing library available. Install PyPDF2 or pdfplumber.")
                    return self._parse_as_text_fallback(file_path)
            
            if pdf_parser == 'pdfplumber':
                content_blocks = self._parse_pdf_with_pdfplumber(file_path)
            else:
                content_blocks = self._parse_pdf_with_pypdf2(file_path)
                
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            # Fallback to basic text parsing
            content_blocks = self._parse_as_text_fallback(file_path)
        
        return content_blocks
    
    def _parse_pdf_with_pdfplumber(self, file_path: Path) -> List[MultimodalContent]:
        """Parse PDF using pdfplumber (better table support)."""
        import pdfplumber
        
        content_blocks = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text
                text = page.extract_text()
                if text and text.strip():
                    # Split text into paragraphs and detect code blocks
                    text_blocks = self._process_text_content(text, f"page_{page_num}")
                    content_blocks.extend(text_blocks)
                
                # Extract tables
                tables = page.extract_tables()
                for table_num, table in enumerate(tables, 1):
                    if table:
                        content_blocks.append(MultimodalContent(
                            content_type='table',
                            content=table,
                            metadata={
                                'page': page_num,
                                'table_index': table_num,
                                'rows': len(table),
                                'columns': len(table[0]) if table else 0
                            },
                            source_location=f"page_{page_num}_table_{table_num}"
                        ))
        
        return content_blocks
    
    def _parse_pdf_with_pypdf2(self, file_path: Path) -> List[MultimodalContent]:
        """Parse PDF using PyPDF2 (basic text extraction)."""
        import PyPDF2
        
        content_blocks = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    text_blocks = self._process_text_content(text, f"page_{page_num}")
                    content_blocks.extend(text_blocks)
        
        return content_blocks
    
    def _parse_docx(self, file_path: Path) -> List[MultimodalContent]:
        """Parse DOCX documents."""
        content_blocks = []
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Parse paragraphs
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                if paragraph.text.strip():
                    text_blocks = self._process_text_content(
                        paragraph.text, 
                        f"paragraph_{para_num}"
                    )
                    content_blocks.extend(text_blocks)
            
            # Parse tables
            for table_num, table in enumerate(doc.tables, 1):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    content_blocks.append(MultimodalContent(
                        content_type='table',
                        content=table_data,
                        metadata={
                            'table_index': table_num,
                            'rows': len(table_data),
                            'columns': len(table_data[0]) if table_data else 0
                        },
                        source_location=f"table_{table_num}"
                    ))
            
        except ImportError:
            logger.warning("python-docx not available. Install with: pip install python-docx")
            content_blocks = self._parse_as_text_fallback(file_path)
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            content_blocks = self._parse_as_text_fallback(file_path)
        
        return content_blocks
    
    def _parse_markdown(self, file_path: Path) -> List[MultimodalContent]:
        """Parse Markdown documents."""
        content_blocks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Split by markdown sections
            sections = re.split(r'\n(?=#{1,6}\s)', content)
            
            for section_num, section in enumerate(sections, 1):
                if section.strip():
                    # Process each section for different content types
                    section_blocks = self._process_markdown_section(
                        section, 
                        f"section_{section_num}"
                    )
                    content_blocks.extend(section_blocks)
                    
        except Exception as e:
            logger.error(f"Error parsing Markdown {file_path}: {e}")
        
        return content_blocks
    
    def _parse_html(self, file_path: Path) -> List[MultimodalContent]:
        """Parse HTML documents."""
        content_blocks = []
        
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Extract text content
            for element_num, element in enumerate(soup.find_all(['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']), 1):
                text = element.get_text().strip()
                if text:
                    text_blocks = self._process_text_content(text, f"element_{element_num}")
                    content_blocks.extend(text_blocks)
            
            # Extract code blocks
            for code_num, code_element in enumerate(soup.find_all(['code', 'pre']), 1):
                code_text = code_element.get_text().strip()
                if code_text:
                    language = code_element.get('class', [''])[0] if code_element.get('class') else 'unknown'
                    content_blocks.append(MultimodalContent(
                        content_type='code',
                        content=code_text,
                        metadata={
                            'language': language,
                            'element_type': code_element.name,
                            'code_index': code_num
                        },
                        source_location=f"code_{code_num}"
                    ))
            
            # Extract tables
            for table_num, table in enumerate(soup.find_all('table'), 1):
                table_data = []
                for row in table.find_all('tr'):
                    row_data = [cell.get_text().strip() for cell in row.find_all(['td', 'th'])]
                    if row_data:
                        table_data.append(row_data)
                
                if table_data:
                    content_blocks.append(MultimodalContent(
                        content_type='table',
                        content=table_data,
                        metadata={
                            'table_index': table_num,
                            'rows': len(table_data),
                            'columns': len(table_data[0]) if table_data else 0
                        },
                        source_location=f"table_{table_num}"
                    ))
            
            # Extract image metadata
            for img_num, img in enumerate(soup.find_all('img'), 1):
                img_src = img.get('src', '')
                img_alt = img.get('alt', '')
                img_title = img.get('title', '')
                
                content_blocks.append(MultimodalContent(
                    content_type='image',
                    content={
                        'src': img_src,
                        'alt': img_alt,
                        'title': img_title,
                        'description': img_alt or img_title or f"Image {img_num}"
                    },
                    metadata={
                        'image_index': img_num,
                        'has_alt': bool(img_alt),
                        'has_title': bool(img_title)
                    },
                    source_location=f"image_{img_num}"
                ))
                
        except ImportError:
            logger.warning("BeautifulSoup not available. Install with: pip install beautifulsoup4")
            content_blocks = self._parse_as_text_fallback(file_path)
        except Exception as e:
            logger.error(f"Error parsing HTML {file_path}: {e}")
            content_blocks = self._parse_as_text_fallback(file_path)
        
        return content_blocks
    
    def _parse_text(self, file_path: Path) -> List[MultimodalContent]:
        """Parse plain text documents."""
        content_blocks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Process as text content
            text_blocks = self._process_text_content(content, "full_document")
            content_blocks.extend(text_blocks)
            
        except Exception as e:
            logger.error(f"Error parsing text file {file_path}: {e}")
        
        return content_blocks
    
    def _process_text_content(self, text: str, location: str) -> List[MultimodalContent]:
        """Process text content with enhanced chunking strategy."""
        content_blocks = []

        try:
            # Use enhanced chunker for better list detection
            from multimodal_processing.enhanced_chunker import EnhancedChunker
            chunker = EnhancedChunker()

            enhanced_chunks = chunker.enhanced_chunk_text(text, location)

            # Convert enhanced chunks to MultimodalContent
            for chunk in enhanced_chunks:
                content_type = self._map_chunk_type_to_content_type(chunk.chunk_type)

                content_blocks.append(MultimodalContent(
                    content_type=content_type,
                    content=chunk.content,
                    metadata={
                        'chunk_type': chunk.chunk_type.value,
                        'priority_score': chunk.priority_score,
                        'list_level': chunk.list_level,
                        'structured_tags': chunk.structured_tags,
                        'word_count': len(chunk.content.split()),
                        'char_count': len(chunk.content),
                        **chunk.metadata
                    },
                    source_location=chunk.source_location
                ))

        except ImportError:
            logger.warning("Enhanced chunker not available, falling back to basic chunking")
            # Fallback to original chunking
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

            for para_num, paragraph in enumerate(paragraphs, 1):
                # Check if this looks like code
                if self._is_code_block(paragraph):
                    language = self._detect_language(paragraph)
                    content_blocks.append(MultimodalContent(
                        content_type='code',
                        content=paragraph,
                        metadata={
                            'language': language,
                            'detected': True,
                            'paragraph_index': para_num
                        },
                        source_location=f"{location}_code_{para_num}"
                    ))
                else:
                    # Regular text
                    content_blocks.append(MultimodalContent(
                        content_type='text',
                        content=paragraph,
                        metadata={
                            'paragraph_index': para_num,
                            'word_count': len(paragraph.split()),
                            'char_count': len(paragraph)
                        },
                        source_location=f"{location}_text_{para_num}"
                    ))

        return content_blocks

    def _map_chunk_type_to_content_type(self, chunk_type) -> str:
        """Map enhanced chunk types to MultimodalContent types."""
        from multimodal_processing.enhanced_chunker import ChunkType

        mapping = {
            ChunkType.NARRATIVE: 'text',
            ChunkType.BULLET_LIST: 'list',
            ChunkType.NUMBERED_LIST: 'list',
            ChunkType.CAPABILITY: 'capability',
            ChunkType.REQUIREMENT: 'requirement',
            ChunkType.HEADER: 'header',
            ChunkType.TABLE_ROW: 'table',
            ChunkType.CODE_BLOCK: 'code',
        }

        return mapping.get(chunk_type, 'text')
    
    def _process_markdown_section(self, section: str, location: str) -> List[MultimodalContent]:
        """Process a markdown section and extract different content types."""
        content_blocks = []
        
        # Extract code blocks (```language ... ```)
        code_pattern = r'```(\w+)?\n(.*?)\n```'
        code_matches = re.finditer(code_pattern, section, re.DOTALL)
        
        for code_num, match in enumerate(code_matches, 1):
            language = match.group(1) or 'unknown'
            code_content = match.group(2).strip()
            
            content_blocks.append(MultimodalContent(
                content_type='code',
                content=code_content,
                metadata={
                    'language': language,
                    'markdown_block': True,
                    'code_index': code_num
                },
                source_location=f"{location}_code_{code_num}"
            ))
        
        # Remove code blocks from text for further processing
        text_without_code = re.sub(code_pattern, '', section, flags=re.DOTALL)
        
        # Extract tables (basic markdown table detection)
        table_pattern = r'\|.*\|\n\|[-\s\|]+\|\n(\|.*\|\n)+'
        table_matches = re.finditer(table_pattern, text_without_code, re.MULTILINE)
        
        for table_num, match in enumerate(table_matches, 1):
            table_text = match.group(0)
            table_data = self._parse_markdown_table(table_text)
            
            if table_data:
                content_blocks.append(MultimodalContent(
                    content_type='table',
                    content=table_data,
                    metadata={
                        'markdown_table': True,
                        'table_index': table_num,
                        'rows': len(table_data),
                        'columns': len(table_data[0]) if table_data else 0
                    },
                    source_location=f"{location}_table_{table_num}"
                ))
        
        # Remove tables from text
        text_without_tables = re.sub(table_pattern, '', text_without_code, flags=re.MULTILINE)
        
        # Process remaining text
        if text_without_tables.strip():
            text_blocks = self._process_text_content(text_without_tables, location)
            content_blocks.extend(text_blocks)
        
        return content_blocks
    
    def _parse_markdown_table(self, table_text: str) -> List[List[str]]:
        """Parse a markdown table into a 2D array."""
        lines = table_text.strip().split('\n')
        table_data = []
        
        for line in lines:
            if '|' in line and not re.match(r'^\|[-\s\|]+\|$', line):
                # This is a data row, not a separator
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                table_data.append(cells)
        
        return table_data
    
    def _is_code_block(self, text: str) -> bool:
        """Detect if a text block is likely code."""
        code_indicators = [
            'def ', 'class ', 'import ', 'from ', 'function',
            '{', '}', '()', '=>', 'var ', 'let ', 'const ',
            'public ', 'private ', 'static ', '#include',
            'SELECT ', 'INSERT ', 'UPDATE ', 'DELETE '
        ]
        
        # Check for code-like patterns
        has_code_keywords = any(indicator in text for indicator in code_indicators)
        has_brackets = '{' in text and '}' in text
        has_semicolons = text.count(';') > 2
        has_indentation = '\n    ' in text or '\n\t' in text
        
        return has_code_keywords or has_brackets or has_semicolons or has_indentation
    
    def _detect_language(self, code: str) -> str:
        """Detect programming language from code content."""
        language_patterns = {
            'python': ['def ', 'import ', 'from ', 'class ', '__init__'],
            'javascript': ['function', 'var ', 'let ', 'const ', '=>'],
            'java': ['public class', 'private ', 'public static'],
            'cpp': ['#include', 'std::', 'cout', 'cin'],
            'sql': ['SELECT', 'INSERT', 'UPDATE', 'DELETE', 'FROM'],
            'bash': ['#!/bin/bash', 'echo ', 'grep ', 'awk '],
            'json': ['{', '":', '}', '[', ']'],
            'yaml': ['---', '- ', ': ']
        }
        
        code_lower = code.lower()
        
        for language, patterns in language_patterns.items():
            if any(pattern.lower() in code_lower for pattern in patterns):
                return language
        
        return 'unknown'

    def _parse_code_file(self, file_path: Path) -> List[MultimodalContent]:
        """Parse code files (Python, JavaScript, Java, C++, etc.)."""
        content_blocks = []

        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            # Detect language from file extension
            language_map = {
                '.py': 'python',
                '.js': 'javascript',
                '.java': 'java',
                '.cpp': 'cpp',
                '.c': 'c'
            }

            language = language_map.get(file_path.suffix.lower(), 'unknown')

            # For now, treat the entire file as one code block
            # Could be enhanced to split into functions/classes
            content_blocks.append(MultimodalContent(
                content_type='code',
                content=content,
                metadata={
                    'language': language,
                    'file_extension': file_path.suffix,
                    'lines_count': len(content.split('\n')),
                    'char_count': len(content),
                    'is_full_file': True
                },
                source_location="full_file"
            ))

        except Exception as e:
            logger.error(f"Error parsing code file {file_path}: {e}")

        return content_blocks

    def _parse_as_text_fallback(self, file_path: Path) -> List[MultimodalContent]:
        """Fallback parser that treats everything as text."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            return [MultimodalContent(
                content_type='text',
                content=content,
                metadata={
                    'fallback_parsing': True,
                    'char_count': len(content),
                    'word_count': len(content.split())
                },
                source_location="full_document"
            )]
        except Exception as e:
            logger.error(f"Fallback parsing failed for {file_path}: {e}")
            return []
    
    def _calculate_parsing_stats(self, content_blocks: List[MultimodalContent]) -> Dict[str, Any]:
        """Calculate statistics about the parsed content."""
        stats = {
            'total_blocks': len(content_blocks),
            'content_types': {},
            'total_text_length': 0,
            'total_code_blocks': 0,
            'total_tables': 0,
            'total_images': 0
        }
        
        for block in content_blocks:
            content_type = block.content_type
            stats['content_types'][content_type] = stats['content_types'].get(content_type, 0) + 1
            
            if content_type == 'text':
                stats['total_text_length'] += len(str(block.content))
            elif content_type == 'code':
                stats['total_code_blocks'] += 1
            elif content_type == 'table':
                stats['total_tables'] += 1
            elif content_type == 'image':
                stats['total_images'] += 1
        
        return stats

# Global parser instance
_document_parser = None

def get_document_parser() -> MultimodalDocumentParser:
    """Get or create a global document parser instance."""
    global _document_parser
    
    if _document_parser is None:
        _document_parser = MultimodalDocumentParser()
    
    return _document_parser
